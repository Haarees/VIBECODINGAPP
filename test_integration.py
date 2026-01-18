import unittest
import os
import shutil
from docx import Document
from app import app, UPLOAD_FOLDER, CONVERTED_FOLDER

class TestConversion(unittest.TestCase):
    def setUp(self):
        app.config['TESTING'] = True
        self.client = app.test_client()
        
        # Create a dummy docx
        self.doc_path = 'test_fixture.docx'
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph for conversion.')
        doc.save(self.doc_path)
        
    def tearDown(self):
        # Cleanup
        if os.path.exists(self.doc_path):
            os.remove(self.doc_path)
        # We might want to keep uploads/converted to see, but for auto-test cleanup is good.
        # Clean specific test files from uploads/converted
        for f in os.listdir(UPLOAD_FOLDER):
            if f.startswith('test_fixture'):
                os.remove(os.path.join(UPLOAD_FOLDER, f))
        for f in os.listdir(CONVERTED_FOLDER):
            if f.startswith('test_fixture'):
                os.remove(os.path.join(CONVERTED_FOLDER, f))

    def test_upload_and_convert(self):
        with open(self.doc_path, 'rb') as f:
            data = {
                'file': (f, 'test_fixture.docx')
            }
            response = self.client.post('/upload', data=data, content_type='multipart/form-data')
            
        self.assertEqual(response.status_code, 200)
        json_data = response.get_json()
        self.assertIn('downloadUrl', json_data)
        
        pdf_url = json_data['downloadUrl']
        self.assertTrue(pdf_url.endswith('.pdf'))
        
        # Verify download link works
        download_response = self.client.get(pdf_url)
        self.assertEqual(download_response.status_code, 200)
        self.assertEqual(download_response.mimetype, 'application/pdf')

if __name__ == '__main__':
    unittest.main()
